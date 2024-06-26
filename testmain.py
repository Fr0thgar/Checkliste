from tkinter import *
import customtkinter
import os
import shutil
import docx
from datetime import datetime

customtkinter.set_default_color_theme("./theme.json")

def show_selected():
    selected_items = [var.get() for var in checkboxes]
    selected_text = [entry.get() for entry in entries]
    name = name_var.get()
    selected_items.append(name)
    selected_items_dict = {
         option: text for option, 
         text in zip(options, selected_text) if selected_items[options.index(option)]}
    
    # Save the updaed document to a new file in a folder named month and date
    current_month = datetime.now().strftime("%B")
    current_date = datetime.now().strftime("%d")
    current_year = datetime.now().strftime("%Y")
    folder_path = os.path.join(os.getcwd(), current_year, current_month, current_date)
    os.makedirs(folder_path, exist_ok=True)
    
    # Specify the file path for the copied template
    file_path = os.path.join(folder_path, f"{name}.docx")

    # Copy the Template.docx file with Employee name
    template_path = r"C:\Users\jbs\Code\Checkliste\Template.docx"
    shutil.copyfile(template_path, file_path)

    # Open the copied template
    doc = docx.Document(file_path)
    for paragraph in doc.paragraphs:
         paragraph.text = "Employee Name: " + name + "\n"

    # Insert the selected items
    doc.add_paragraph("Selected Items: \n")
    for option, text in selected_items_dict.items():
         doc.add_paragraph(f"{option}: {text}\n")

    # Save the updated document 
    doc.save(file_path)
    
    # Clear the checkmarks, text, and name input
    for var in checkboxes:
        var.set(0)
    for entry in entries:
         entry.delete(0, customtkinter.END)
    name_var.set("")
    
    """# Creates the Docx file
    doc = docx.Document()
    doc.add_paragraph("Employee Name: " + name + "\n")
    doc.add_paragraph("Selected Items: \n")
    for option in selected_items_dict.items():
         doc.add_paragraph(f"{option}: {text}: \n")
    doc.save(os.path.join(folder_path, name + ".docx"))"""

    """# Commented out section 
    #with open(file_path, "w") as f:
    #    f.write("Employee Name: " + name + "\n")
    #    f.write("Selected items:\n")
    #   for option, text in selected_items_dict.items():
    #            f.write(f"{option}: {text}\n")
    #print("Selected items have been saved to " + name + ".txt")"""
    
# Create the main application window
root = customtkinter.CTk()
root.title("Checkbox Example")
root.geometry("500x500")

#Create a label for the name input
name_label = customtkinter.CTkLabel(root, text="Employee name:")
name_label.pack()

#Create a textbox for name input
name_var = customtkinter.StringVar()
name_entry = customtkinter.CTkEntry(root, textvariable=name_var)
name_entry.pack()

# Create some checkboxes
options = ["Computer", "Mobil", "Headset", "Nøgle Brik", 
           "Oplader", "Taske", "Andet"]
checkboxes = []
entries = []
for option in options:
    var = customtkinter.IntVar()
    checkbox = customtkinter.CTkCheckBox(root, text=option, variable=var)
    checkbox.pack(anchor=customtkinter.W)
    checkboxes.append(var)
    entry = customtkinter.CTkEntry(root)
    entry.pack()
    entries.append(entry)

# Create a button to submit selected checkboxes
show_button = customtkinter.CTkButton(master = root, text="Submit", 
                                      command=show_selected)
show_button.pack()

# Start the GUI event loop
root.mainloop()

